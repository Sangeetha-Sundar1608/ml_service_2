"""
DOCX Processor - Extract text from Word documents
"""

import logging
from docx import Document
from typing import Dict, Any

logger = logging.getLogger(__name__)


class DocxExtractionResult:
    def __init__(self, text: str, format: str = "docx"):
        self.text = text
        self.format = format

    def dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "format": self.format
        }


class DOCXProcessor:
    async def extract(self, file_path: str) -> DocxExtractionResult:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            full_text = "\n".join(paragraphs)
            
            return DocxExtractionResult(
                text=full_text,
                format="docx"
            )
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
